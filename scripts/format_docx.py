"""
Post-process pandoc-generated docx to match 研电赛 template formatting.

Template requirements:
- Page: A4, margins top/bottom 2.54cm, left/right 3.18cm
- Body: 宋体 小四号 12pt, line spacing fixed 20pt
- Chapter title (第N章): 黑体 小二号 18pt, centered
- Section (N.N): 黑体 小三号 15pt, left-aligned
- Subsection (N.N.N): 黑体 小四号 12pt, left-aligned
- Header: 宋体 五号 10.5pt, centered, two lines
- Footer: page number centered
"""
import re
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT = "研电赛论文_pandoc.docx"
OUTPUT = "研电赛论文.docx"

SONG_TI = "宋体"
HEI_TI = "黑体"

# Chinese font names with fallbacks
SONG_FONTS = ["宋体", "SimSun", "STSong"]
HEI_FONTS = ["黑体", "SimHei", "STHeiti"]


def set_run_font(run, font_name, size_pt, bold=False):
    """Set font for a run, including CJK font."""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_paragraph_spacing(para, line_spacing_pt=20, before_pt=0, after_pt=0):
    """Set fixed line spacing."""
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} />')
        pPr.append(spacing)
    # Fixed line spacing (line value in twips: 1pt = 20 twips)
    spacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    if before_pt:
        spacing.set(qn('w:before'), str(int(before_pt * 20)))
    if after_pt:
        spacing.set(qn('w:after'), str(int(after_pt * 20)))


def add_header(section, line1, line2, font_size=10.5):
    """Add two-line centered header."""
    header = section.header
    header.is_linked_to_previous = False
    # Clear existing
    for p in header.paragraphs:
        p.clear()
    # First line
    p1 = header.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(line1)
    set_run_font(run1, SONG_TI, font_size)
    # Second line
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(line2)
    set_run_font(run2, SONG_TI, font_size)


def add_page_number(section):
    """Add centered page number to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    # Add page number field
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def is_chapter_heading(text):
    """Check if paragraph is a chapter heading."""
    # Matches patterns like "作品难点与创新", "方案论证与设计", etc.
    chapter_titles = [
        "作品难点与创新", "方案论证与设计", "原理分析与硬件电路图",
        "软件设计与流程", "系统测试与分析", "总结",
        "摘要", "ABSTRACT", "目录", "参考文献"
    ]
    stripped = text.strip()
    return any(t in stripped for t in chapter_titles)


def detect_heading_level(para):
    """Detect heading level from style or content."""
    style_name = para.style.name if para.style else ""
    if "Heading 1" in style_name:
        return 1
    if "Heading 2" in style_name:
        return 2
    if "Heading 3" in style_name:
        return 3
    return 0


def main():
    doc = Document(INPUT)

    # --- Page setup ---
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        section.orientation = WD_ORIENT.PORTRAIT

        # Add header
        add_header(
            section,
            "第二十一届中国研究生电子设计竞赛",
            "面向eZ-PLM的电子元器件智能选型与风险评估Agent系统"
        )
        # Add page number
        add_page_number(section)

    # --- Format paragraphs ---
    for para in doc.paragraphs:
        level = detect_heading_level(para)
        text = para.text.strip()

        if not text:
            continue

        if level == 1 or is_chapter_heading(text):
            # Chapter title: 黑体 小二号 18pt, centered
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(para, line_spacing_pt=20, before_pt=0, after_pt=20)
            for run in para.runs:
                set_run_font(run, HEI_TI, 18, bold=True)

        elif level == 2:
            # Section: 黑体 小三号 15pt, left
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(para, line_spacing_pt=20, before_pt=20, after_pt=10)
            for run in para.runs:
                set_run_font(run, HEI_TI, 15, bold=True)

        elif level == 3:
            # Subsection: 黑体 小四号 12pt, left
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(para, line_spacing_pt=20, before_pt=15, after_pt=6)
            for run in para.runs:
                set_run_font(run, HEI_TI, 12, bold=True)

        else:
            # Body text: 宋体 小四号 12pt
            set_paragraph_spacing(para, line_spacing_pt=20)
            for run in para.runs:
                set_run_font(run, SONG_TI, 12)

    # --- Format tables ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    set_paragraph_spacing(para, line_spacing_pt=20)
                    for run in para.runs:
                        set_run_font(run, SONG_TI, 10.5)

    doc.save(OUTPUT)
    print(f"Saved formatted docx: {OUTPUT}")

    # Count characters
    total_cn = 0
    for para in Document(OUTPUT).paragraphs:
        total_cn += len(re.findall(r'[\u4e00-\u9fff]', para.text))
    print(f"Chinese characters (body text): {total_cn}")


if __name__ == "__main__":
    main()
